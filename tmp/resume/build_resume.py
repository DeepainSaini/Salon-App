from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\full stack\salon app")
OUTPUT = ROOT / "output" / "docx" / "Deepain_Saini_Google_SAD_Apprenticeship_Resume.docx"

FONT = "Arial"
BLACK = RGBColor(0x18, 0x18, 0x18)
MUTED = RGBColor(0x4A, 0x4A, 0x4A)
ACCENT = RGBColor(0x17, 0x4E, 0x7A)
LINK_BLUE = "1155CC"


def set_cell_free_font(run, size, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return run


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_hyperlink(paragraph, text, url, size=9.8, bold=False, color=LINK_BLUE):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    properties.append(fonts)

    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    properties.append(run_color)

    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(size * 2)))
    properties.append(size_element)

    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(int(size * 2)))
    properties.append(size_cs)

    if bold:
        properties.append(OxmlElement("w:b"))

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(properties)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(15)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    set_cell_free_font(run, 11.7, bold=True, color=ACCENT)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.first_line_indent = Inches(-0.14)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_cell_free_font(run, 9.8, color=BLACK)
    return paragraph


def add_project(doc, title, url, dates, stack, bullets):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(11)
    heading.paragraph_format.space_after = Pt(2.5)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
    add_hyperlink(heading, title, url, size=10.45, bold=True, color=LINK_BLUE)
    date_run = heading.add_run(f"\t{dates}")
    set_cell_free_font(date_run, 9.5, bold=True, color=MUTED)

    stack_paragraph = doc.add_paragraph()
    stack_paragraph.paragraph_format.space_before = Pt(0)
    stack_paragraph.paragraph_format.space_after = Pt(4.5)
    stack_paragraph.paragraph_format.keep_with_next = True
    stack_run = stack_paragraph.add_run(stack)
    set_cell_free_font(stack_run, 9.15, italic=True, color=MUTED)

    for bullet in bullets:
        add_bullet(doc, bullet)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.8)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    heading = styles["Heading 1"]
    heading.font.name = FONT
    heading._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    heading._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    heading.font.size = Pt(11.7)
    heading.font.bold = True
    heading.font.color.rgb = ACCENT
    heading.paragraph_format.space_before = Pt(15)
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True

    list_bullet = styles["List Bullet"]
    list_bullet.font.name = FONT
    list_bullet._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    list_bullet._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    list_bullet.font.size = Pt(9.8)
    list_bullet.font.color.rgb = BLACK
    list_bullet.paragraph_format.left_indent = Inches(0.22)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.14)
    list_bullet.paragraph_format.space_after = Pt(8)
    list_bullet.paragraph_format.line_spacing = 1.12

    if "Contact Line" not in styles:
        contact = styles.add_style("Contact Line", WD_STYLE_TYPE.PARAGRAPH)
    else:
        contact = styles["Contact Line"]
    contact.font.name = FONT
    contact.font.size = Pt(9.8)
    contact.font.color.rgb = MUTED
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(9)
    contact.paragraph_format.line_spacing = 1.0


def build_resume():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_before = Pt(0)
    name.paragraph_format.space_after = Pt(3.5)
    name.paragraph_format.keep_with_next = True
    set_cell_free_font(name.add_run("DEEPAIN SAINI"), 21.5, bold=True, color=BLACK)

    contact = doc.add_paragraph(style="Contact Line")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(contact, "deepainsaini111@gmail.com", "mailto:deepainsaini111@gmail.com", size=9.7)
    set_cell_free_font(contact.add_run("  |  +91 96499 90222  |  "), 9.7, color=MUTED)
    add_hyperlink(contact, "github.com/DeepainSaini", "https://github.com/DeepainSaini", size=9.7)

    add_section_heading(doc, "Professional Summary")
    summary = doc.add_paragraph()
    summary.paragraph_format.space_after = Pt(6)
    summary.paragraph_format.line_spacing = 1.08
    set_cell_free_font(
        summary.add_run(
            "Entry-level software developer with hands-on experience building and deploying full-stack personal projects using JavaScript, Node.js, Express.js, MySQL, and Sequelize. Developed REST APIs, authentication, database transactions, real-time communication, cloud integrations, background jobs, and API validation. Comfortable collaborating with teammates through Git and GitHub branch, pull, push, and merge workflows, and eager to contribute to implementation, integration, testing, and high-quality software through Google's Software Application Development Apprenticeship."
        ),
        9.8,
    )

    add_section_heading(doc, "Education")
    education = doc.add_paragraph()
    education.paragraph_format.space_after = Pt(2.5)
    education.paragraph_format.keep_with_next = True
    set_cell_free_font(education.add_run("Bachelor of Technology in Computer Science and Engineering"), 10.05, bold=True)
    set_cell_free_font(education.add_run("  |  Graduated: 2021"), 9.8, bold=True, color=MUTED)
    university = doc.add_paragraph()
    university.paragraph_format.space_after = Pt(4.5)
    set_cell_free_font(
        university.add_run("Guru Jambheshwar University of Science and Technology, Hisar"),
        9.65,
        color=MUTED,
    )

    add_section_heading(doc, "Technical Skills")
    skill_lines = [
        ("Programming", "JavaScript (ES6+), SQL, HTML, CSS"),
        ("Backend", "Node.js, Express.js, REST APIs, Socket.IO, JWT, bcrypt, scheduled jobs"),
        ("Databases", "MySQL, Sequelize ORM, migrations, transactions, pagination, indexing"),
        ("Testing and Quality", "Postman API testing, input validation, error handling, Winston logging"),
        ("Cloud and Tools", "AWS S3, AWS RDS, AWS EC2, Nginx, PM2, Jenkins CI/CD"),
        ("Productivity", "Google Workspace: Gmail, Drive, Docs, Sheets, Slides, Calendar, Meet"),
        ("Collaboration", "Feature decomposition, task planning, teammate coordination, and Git/GitHub branch and merge workflows"),
    ]
    for label, value in skill_lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3.5)
        set_cell_free_font(paragraph.add_run(f"{label}: "), 9.6, bold=True)
        set_cell_free_font(paragraph.add_run(value), 9.6)

    add_section_heading(doc, "Software Development Projects")

    add_project(
        doc,
        "Salon Appointment Platform",
        "https://github.com/DeepainSaini/Salon-App",
        "Jul 2026 - Present",
        "JavaScript, Node.js, Express.js, MySQL, Sequelize, JWT, Cashfree, Brevo",
        [
            "Planned and built customer, staff, and admin workflows by breaking booking requirements into dynamic slot calculation, availability and conflict checks, rescheduling, cancellation, and reviews.",
            "Integrated JWT/cookie authentication, Cashfree sandbox payments, Brevo confirmation and reminder emails, and Sequelize transactions for multi-step appointment workflows.",
        ],
    )

    add_project(
        doc,
        "Group Chat Application",
        "https://github.com/DeepainSaini/group-chat-app",
        "Sep 2025 - May 2026",
        "JavaScript, Node.js, Express.js, Socket.IO, MySQL, Sequelize, AWS S3",
        [
            "Built authenticated personal and group messaging with Socket.IO rooms, persistent MySQL storage, and server-side membership checks before group-message delivery.",
            "Added AWS S3 media sharing and a scheduled archival job that moves messages older than 24 hours to a separate table to keep active-message queries focused.",
        ],
    )

    add_project(
        doc,
        "Daily Expense Tracker",
        "https://github.com/DeepainSaini/daily-expense-tracking",
        "May 2025 - Sep 2025",
        "JavaScript, Node.js, Express.js, MySQL, Sequelize, AWS, Cashfree, JWT",
        [
            "Developed expense tracking with JWT/bcrypt authentication, pagination, transactional create/delete flows, password-reset emails, premium payments, and AWS S3 report exports.",
            "Deployed on AWS EC2 with MySQL on RDS, PM2, Jenkins CI/CD, and Winston logging; optimized leaderboard reads by maintaining per-user expense totals.",
        ],
    )

    doc.core_properties.title = "Deepain Saini - Google Software Application Development Apprenticeship Resume"
    doc.core_properties.subject = "Application for Software Application Development Apprenticeship, March 2027 Start"
    doc.core_properties.author = "Deepain Saini"
    doc.core_properties.keywords = (
        "software application development, JavaScript, Node.js, Express.js, MySQL, Sequelize, REST APIs, "
        "Socket.IO, implementation, integration, problem solving, Git, Postman"
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_resume()
