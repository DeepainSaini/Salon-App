import json
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document
from pypdf import PdfReader


ROOT = Path(r"D:\full stack\salon app")
DOCX = ROOT / "output" / "docx" / "Deepain_Saini_Google_SAD_Apprenticeship_Resume.docx"
PDF = ROOT / "output" / "pdf" / "Deepain_Saini_Google_SAD_Apprenticeship_Resume.pdf"

EXPECTED_URLS = {
    "mailto:deepainsaini111@gmail.com",
    "https://github.com/DeepainSaini",
    "https://github.com/DeepainSaini/Salon-App",
    "https://github.com/DeepainSaini/group-chat-app",
    "https://github.com/DeepainSaini/daily-expense-tracking",
}

REQUIRED_TEXT = [
    "DEEPAIN SAINI",
    "PROFESSIONAL SUMMARY",
    "EDUCATION",
    "Graduated: 2021",
    "TECHNICAL SKILLS",
    "Testing and Quality",
    "Nginx",
    "Google Workspace",
    "Collaboration",
    "Feature decomposition",
    "task planning",
    "SOFTWARE DEVELOPMENT PROJECTS",
    "Salon Appointment Platform",
    "Group Chat Application",
    "Daily Expense Tracker",
    "implementation",
    "integration",
    "Postman API testing",
]


def verify():
    assert DOCX.exists() and DOCX.stat().st_size > 10_000, "DOCX is missing or unexpectedly small"
    assert PDF.exists() and PDF.stat().st_size > 20_000, "PDF is missing or unexpectedly small"

    document = Document(DOCX)
    assert len(document.sections) == 1, "DOCX must contain exactly one section"
    section = document.sections[0]
    assert abs(section.page_width.inches - 8.5) < 0.01, "DOCX width is not US Letter"
    assert abs(section.page_height.inches - 11.0) < 0.01, "DOCX height is not US Letter"
    assert len(document.tables) == 0, "ATS resume should not contain tables"
    assert len(document.inline_shapes) == 0, "ATS resume should not contain images or inline shapes"

    doc_text = "\n".join(p.text for p in document.paragraphs)
    for value in REQUIRED_TEXT:
        assert value in doc_text, f"Required DOCX text missing: {value}"
    assert "TODO" not in doc_text and "TBD" not in doc_text, "Placeholder text found"
    assert "codex-file-citation" not in doc_text, "Internal citation token found"

    docx_urls = {
        rel.target_ref
        for rel in document.part.rels.values()
        if rel.reltype.endswith("/hyperlink")
    }
    assert EXPECTED_URLS.issubset(docx_urls), f"DOCX hyperlink(s) missing: {sorted(EXPECTED_URLS - docx_urls)}"

    with zipfile.ZipFile(DOCX) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        assert "<w:tbl" not in xml, "Table XML found"
        assert "w:txbxContent" not in xml, "Text box XML found"

    reader = PdfReader(PDF)
    assert len(reader.pages) == 1, "PDF must contain exactly one page"
    page = reader.pages[0]
    width = float(page.mediabox.width)
    height = float(page.mediabox.height)
    assert abs(width - 612) < 1 and abs(height - 792) < 1, "PDF is not US Letter"

    pdf_text = page.extract_text() or ""
    for value in REQUIRED_TEXT:
        assert value in pdf_text, f"Required PDF text missing: {value}"

    ordered_sections = [
        "PROFESSIONAL SUMMARY",
        "EDUCATION",
        "TECHNICAL SKILLS",
        "SOFTWARE DEVELOPMENT PROJECTS",
        "Salon Appointment Platform",
        "Group Chat Application",
        "Daily Expense Tracker",
    ]
    positions = [pdf_text.index(value) for value in ordered_sections]
    assert positions == sorted(positions), "PDF text extraction order is not ATS-friendly"

    pdf_urls = set()
    for annotation_ref in page.get("/Annots", []):
        annotation = annotation_ref.get_object()
        action = annotation.get("/A")
        if action and action.get("/URI"):
            pdf_urls.add(str(action.get("/URI")))
    assert EXPECTED_URLS.issubset(pdf_urls), f"PDF hyperlink(s) missing: {sorted(EXPECTED_URLS - pdf_urls)}"

    with pdfplumber.open(PDF) as opened:
        plumber_page = opened.pages[0]
        words = plumber_page.extract_words()
        assert words, "PDF contains no extractable text"
        last_bottom = max(word["bottom"] for word in words)
        bottom_margin = plumber_page.height - last_bottom
        assert bottom_margin > 36, f"PDF content is too close to the bottom edge: {bottom_margin:.1f} pt"

    disallowed_dashes = {"\u2010", "\u2011", "\u2012", "\u2013", "\u2014"}
    assert not any(dash in pdf_text for dash in disallowed_dashes), "Non-ASCII dash found in PDF"

    result = {
        "docx_bytes": DOCX.stat().st_size,
        "pdf_bytes": PDF.stat().st_size,
        "pdf_pages": len(reader.pages),
        "page_size_points": [width, height],
        "docx_hyperlinks": len(docx_urls),
        "pdf_hyperlinks": len(pdf_urls),
        "bottom_margin_points": round(bottom_margin, 1),
        "ats_text_order": "pass",
        "required_content": "pass",
        "single_column_no_tables_or_text_boxes": "pass",
    }
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    verify()
